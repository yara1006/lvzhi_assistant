"""
合同下载接口（独立文件，不影响原有代码）
"""
from typing import Annotated
from io import BytesIO

from fastapi import APIRouter, Depends, Header, Query
from fastapi.responses import StreamingResponse
from sqlalchemy import select
from sqlalchemy.ext.asyncio import AsyncSession
from docx import Document

from app.api.deps import require_service_auth, resolve_user_id
from app.core.config import Settings, get_settings
from app.core.exceptions import AppError
from app.db.models.contract import Contract
from app.db.session import get_db
from app.services.users import get_or_create_user

router = APIRouter(prefix="/contracts", tags=["contracts"])


@router.get(
    "/download/{contract_id}",
    dependencies=[Depends(require_service_auth)],
)
async def download_contract(
    contract_id: str,
    session: Annotated[AsyncSession, Depends(get_db)],
    settings: Annotated[Settings, Depends(get_settings)],
    authorization: Annotated[str | None, Header()] = None,
    user_id: str | None = Query(None, description="用户ID，与生成合同时使用的保持一致"),
):
    """
    下载合同为 Word 文档
    
    - **contract_id**: 合同生成接口返回的合同ID
    - **user_id**: 生成合同时使用的 user_id（可选，如果不提供则从认证信息中解析）
    """
    # 解析用户ID：优先使用 query 参数中的 user_id
    if user_id:
        uid = user_id
    else:
        uid = resolve_user_id(settings, authorization, None)
    
    user = await get_or_create_user(session, uid, settings)
    
    # 查询合同（验证归属）
    result = await session.execute(
        select(Contract).where(
            Contract.id == contract_id,
            Contract.user_id == user.id,
        )
    )
    contract = result.scalar_one_or_none()
    
    if not contract:
        raise AppError("not_found", "合同不存在", status_code=404)
    
    # 创建 Word 文档
    doc = Document()
    
    title = contract.title or contract.contract_type or "合同"
    doc.add_heading(title, 0)
    
    doc.add_heading("基本信息", level=1)
    doc.add_paragraph(f"合同ID：{contract.id}")
    if contract.contract_type:
        doc.add_paragraph(f"合同类型：{contract.contract_type}")
    
    if contract.content:
        doc.add_heading("合同正文", level=1)
        paragraphs = contract.content.split('\n')
        for para in paragraphs:
            if para.strip():
                doc.add_paragraph(para.strip())
    
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    
    return StreamingResponse(
        buffer,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f"attachment; filename=contract_{contract_id}.docx"}
    )
